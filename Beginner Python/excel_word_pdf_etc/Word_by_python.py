import docx

d=docx.Document('report.docx')
print(d.paragraphs) #this will return list of paragraph object
print(d.paragraphs[0])
print(d.paragraphs[3].text)

#Every time style or font or size changes till that we call it as run here
p=d.paragraphs[140]
print(p.runs)
print(p.runs[0].text) #the text will print after which pattern is changes in 140th paragraph of document
print(p.runs[0].bold) #it will print True if the Text is bold can also be done for italic and Underline if not bold will print None

#now we will change the content from program
p.runs[0].bold==None
p.runs[0].italic=True
p.runs[0].underline=True  #this statement will make that written part italic and underlined and not bold
p.runs[0].text = 'italic and Underline text Created ' #it will actually change Data with above mentioned Changes

d.save('report.docx') #Saving Changes to Document you can change Name Also
p=d.paragraphs[141]
#p.style = 'bold' #will make given paragrapgh in bold style



#creating New Docx document

d1=docx.Document()
d1.add_paragraph('hello This is a paragrapgh created by python program')
d1.add_paragraph('This is another paragrapgh')
d1.save('wordByPython.docx')

p=d1.paragraphs[0]
p.add_run('This is the suffix added after creating file')
p.runs[0].bold=True
d1.save('wordByPython.docx') #saving changes